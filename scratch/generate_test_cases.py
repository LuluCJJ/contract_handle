# -*- coding: utf-8 -*-
"""
V2 - 高质量测试用例生成：像真实业务员一样全面填写每个模板的所有字段。
基于 deep_analyze_template.py 的精确坐标映射。
"""
import os, sys, json, shutil
from pathlib import Path
from docx import Document

sys.path.append(os.getcwd())

PROJECT_ROOT = Path(r"d:\AI\project\contract_handle")
TEMPLATE_DIR = PROJECT_ROOT / "inputs" / "bank_template"
CONVERTED_DIR = PROJECT_ROOT / "_converted_docx"
TEST_DATA_DIR = PROJECT_ROOT / "test_data"

# ---- 找到可用的证件图片 ----
EXISTING_ID_IMGS = []
for case_dir in sorted(TEST_DATA_DIR.iterdir()):
    if case_dir.is_dir():
        for f in case_dir.iterdir():
            if f.suffix in (".jpg", ".jpeg", ".png"):
                EXISTING_ID_IMGS.append(str(f))
                break
ID_IMG_1 = EXISTING_ID_IMGS[0] if EXISTING_ID_IMGS else None
ID_IMG_2 = EXISTING_ID_IMGS[1] if len(EXISTING_ID_IMGS) > 1 else ID_IMG_1


def fill_docx_cells(template_path, field_map, output_path):
    """
    Fill cells in a docx template.
    field_map: dict of (table_idx, row_idx, col_idx) -> value
    
    Important: For merged cells, we write to the FIRST column index of the merge span.
    The value will appear across the entire merged range.
    """
    doc = Document(template_path)
    for (ti, ri, ci), value in field_map.items():
        if value is None:
            continue
        value = str(value)
        if ti < len(doc.tables):
            table = doc.tables[ti]
            if ri < len(table.rows):
                row = table.rows[ri]
                if ci < len(row.cells):
                    cell = row.cells[ci]
                    # Clear existing content and write new value
                    for p in cell.paragraphs:
                        p.text = ""
                    cell.paragraphs[0].text = value
    try:
        doc.save(output_path)
    except PermissionError:
        # File locked by another process - save with alt name
        alt_path = output_path.replace(".docx", "_v2.docx")
        doc.save(alt_path)
        print(f"  [WARN] File locked, saved as: {alt_path}")
        return
    print(f"  [OK] Saved: {output_path}")


def create_case(case_id, description, template_src, field_map, eflow_data, id_img_src):
    """Create a complete test case directory."""
    case_dir = TEST_DATA_DIR / case_id
    # Overwrite existing files (don't delete dir to avoid file lock issues)
    case_dir.mkdir(parents=True, exist_ok=True)
    
    # README
    (case_dir / "README.md").write_text(
        f"# {case_id}\n\n**场景**: {description}\n", encoding="utf-8")
    
    # Bank app document
    src_ext = Path(template_src).suffix.lower()
    if src_ext == ".pdf":
        shutil.copy2(template_src, case_dir / "bank_app.pdf")
    else:
        fill_docx_cells(template_src, field_map, str(case_dir / "bank_app.docx"))
    
    # E-Flow JSON
    (case_dir / "eflow.json").write_text(
        json.dumps(eflow_data, ensure_ascii=False, indent=2), encoding="utf-8")
    
    # ID image
    if id_img_src and Path(id_img_src).exists():
        shutil.copy2(id_img_src, case_dir / f"id_document{Path(id_img_src).suffix}")
    
    print(f"  [DONE] Case {case_id}")


# ======================================================================
# Template paths
# ======================================================================
CCB = str(CONVERTED_DIR / "[perfect]ccb_corp_app.docx")
BOC_FRANCE = str(CONVERTED_DIR / "[1]boc_corp_app.docx")
BOC_DOMESTIC = str(CONVERTED_DIR / "中国银行股份有限公司网上银行服务申请变更表(表1.企业客户基本信息表).docx")
BOC_HK = str(TEMPLATE_DIR / "中国银行香港表1.docx")
ICBC = str(CONVERTED_DIR / "ICBC企业客户证书及分支机构信息表.docx")


# ######################################################################
# CASE 010: BOC France e-Corp — PASS (全面填写)
# Template: [1]boc_corp_app.docx
# Table 2 (5x4): 公司基本信息 — Label in C0, Value in C1-3
# Table 3 (16x11): 产品服务 & 账户信息 — 服务类型勾选 + 账号填写
# Table 4 (6x16): 授权模板信息
# Table 8 (21x5): 操作员信息 — Label in C0, Value in C1-4
# ######################################################################
print("\n=== Generating case_010_boc_france_pass ===")
create_case(
    case_id="case_010_boc_france_pass",
    description="[正例] BOC法国 e-Corp 企业网银 - 中英双语全量填写，信息完全一致",
    template_src=BOC_FRANCE,
    field_map={
        # --- Table 2: Company Info (5 rows x 4 cols) ---
        # R0: Company Name -> C1
        (2, 0, 1): "Global Trading Solutions SARL",
        # R1: Account Number(First 6 digits) -> C1
        (2, 1, 1): "880012",
        # R2: Company Address -> C1
        (2, 2, 1): "15 Rue de la Paix, 75002 Paris, France",
        # R3: Email Address -> C1
        (2, 3, 1): "contact@globaltrading.fr",
        # R4: Phone -> C1, Fax -> C3
        (2, 4, 1): "+33 1 42 00 1234",
        (2, 4, 3): "+33 1 42 00 1235",

        # --- Table 3: Account Info (R6-R15 are data rows) ---
        # R6 first account row: Account Number -> C0, Currency -> C4,
        # Single TR Limit -> C5, Daily Limit -> C7, Auth Model -> C9
        (3, 6, 0): "FR7630004000031234567890143",
        (3, 6, 4): "EUR",
        (3, 6, 5): "50,000.00",
        (3, 6, 7): "200,000.00",
        (3, 6, 9): "Model_A",
        # R7 second account
        (3, 7, 0): "FR7630004000039876543210298",
        (3, 7, 4): "USD",
        (3, 7, 5): "100,000.00",
        (3, 7, 7): "500,000.00",
        (3, 7, 9): "Model_A",

        # --- Table 4: Authorisation Info ---
        # R1: Model One Name -> C3
        (4, 1, 3): "Standard Dual Auth",
        # R3: Level I amount -> C1, person count -> C2, operator -> C4
        (4, 3, 1): "50,000 EUR",
        (4, 3, 2): "1",
        (4, 3, 4): "Li Wei / User001",
        # R4: Level II
        (4, 4, 1): "200,000 EUR",
        (4, 4, 2): "2",
        (4, 4, 4): "Li Wei + Chen Fang / User001+User002",

        # --- Table 8: User 1 Operator Info (21 rows x 5 cols) ---
        # R1: Name -> C1
        (8, 1, 1): "Li Wei",
        # R2: Company Name -> C1
        (8, 2, 1): "Global Trading Solutions SARL",
        # R3: ID Type -> C1, ID Number -> C3
        (8, 3, 1): "Passport",
        (8, 3, 3): "E12345678",
        # R4: Email -> C1
        (8, 4, 1): "liwei@globaltrading.fr",
        # R5: Fax -> C1, Phone -> C3
        (8, 5, 1): "+33 1 42 00 1235",
        (8, 5, 3): "+33 1 42 00 1234",
        # R6: Mobile -> C1
        (8, 6, 1): "+33 6 12 34 56 78",
        # R10-R11: Account access rows
        (8, 10, 0): "FR7630004000031234567890143",
        (8, 11, 0): "FR7630004000039876543210298",
    },
    eflow_data={
        "flow_id": "EF2026041001",
        "company": {
            "name": "Global Trading Solutions SARL",
            "name_en": "Global Trading Solutions SARL",
            "cert_type": "Business Registration",
            "cert_number": "RCS Paris B 880 012 345"
        },
        "account": {
            "bank_name": "Bank of China Paris Branch",
            "branch": "Paris Main",
            "account_number": "FR7630004000031234567890143"
        },
        "operator": {"name": "Li Wei", "id_type": "Passport", "id_number": "E12345678"},
        "handler": {"name": "Li Wei", "id_type": "Passport", "id_number": "E12345678"},
        "activity": "Apply for e-Corp Internet Banking Service",
        "permissions": {"level": "Level_A", "single_limit": 50000, "daily_limit": 200000},
        "apply_date": "2026-04-10"
    },
    id_img_src=ID_IMG_1
)


# ######################################################################
# CASE 011: BOC France — MISMATCH (证件号不一致)
# ######################################################################
print("\n=== Generating case_011_boc_france_mismatch ===")
create_case(
    case_id="case_011_boc_france_mismatch",
    description="[负例] BOC法国 - 操作员证件号在Word(E12345678)与eflow(E99999999)中不一致 (CRITICAL)",
    template_src=BOC_FRANCE,
    field_map={
        (2, 0, 1): "Global Trading Solutions SARL",
        (2, 1, 1): "880012",
        (2, 2, 1): "15 Rue de la Paix, 75002 Paris, France",
        (2, 3, 1): "contact@globaltrading.fr",
        (2, 4, 1): "+33 1 42 00 1234",
        (2, 4, 3): "+33 1 42 00 1235",
        (3, 6, 0): "FR7630004000031234567890143",
        (3, 6, 4): "EUR",
        (3, 6, 5): "50,000.00",
        (3, 6, 7): "200,000.00",
        (8, 1, 1): "Li Wei",
        (8, 2, 1): "Global Trading Solutions SARL",
        (8, 3, 1): "Passport",
        (8, 3, 3): "E12345678",       # Word says E12345678
        (8, 4, 1): "liwei@globaltrading.fr",
        (8, 6, 1): "+33 6 12 34 56 78",
    },
    eflow_data={
        "flow_id": "EF2026041101",
        "company": {"name": "Global Trading Solutions SARL", "cert_type": "Business Registration", "cert_number": "RCS Paris B 880 012 345"},
        "account": {"bank_name": "Bank of China Paris Branch", "branch": "Paris Main", "account_number": "FR7630004000031234567890143"},
        "operator": {"name": "Li Wei", "id_type": "Passport", "id_number": "E99999999"},  # MISMATCH!
        "handler": {"name": "Li Wei", "id_type": "Passport", "id_number": "E99999999"},
        "activity": "Apply for e-Corp Internet Banking Service",
        "permissions": {"level": "Level_A", "single_limit": 50000, "daily_limit": 200000},
        "apply_date": "2026-04-11"
    },
    id_img_src=ID_IMG_2
)


# ######################################################################
# CASE 014: BOC Domestic — PASS (国内中行 43x42 超大表格全面填写)
# Template: 中国银行...表1.docx — Single table 0 (43 rows x 42 cols)
# 基于 deep analysis:
#   R0: 申请/服务/类型 C0-2=label, C3-6=新增, C7-41=服务内容(勾选区)
#   R3: *申请单位名称 C0-6=label, C7-41=value
#   R4: *证件类别 C0-6=label, C7-20=checkbox, C21-23=*证件号码, C24-29=value, C30-36=*组织机构代码, C37-41=value
#   R5: *法定代表人 C0-6=label, C7-20=value, C21-23=单位电话, C24-29=value, C30-36=邮政编码, C37-41=value
#   R6: 通信地址 C0-6=label, C7-23=value, C24-29=单位E-MAIL地址, C30-41=value
#   R7: *市场细分 C0-6=label, C7-41=checkbox
#   R8-R9: 服务选择 WEB渠道/对接渠道 (checkbox)
#   R10: 所属集团名称 C0-6=label, C7-25=value, C26-31=集团受理单位, C32-41=value
#   R14-R19: 客户账户信息
#   R21-R26: 操作员信息
# ######################################################################
print("\n=== Generating case_014_boc_domestic_pass ===")
create_case(
    case_id="case_014_boc_domestic_pass",
    description="[正例] 国内中行企业网银申请变更表(43x42超大表格) - 像真实业务员一样全面填写",
    template_src=BOC_DOMESTIC,
    field_map={
        # R0: 服务类型 — 勾选"新增"下的"网银新开户"
        (0, 0, 7): "☑网银新开户（包括：客户注册、服务维护、账户注册、账户限额修改、操作员注册、认证工具申请/绑定等一系列新增服务类型）",
        # R3: *申请单位名称
        (0, 3, 7): "上海星辰国际贸易有限公司",
        # R4: *证件类别 (勾选营业执照), *证件号码, *组织机构代码
        (0, 4, 7): "☑营业执照",
        (0, 4, 24): "91310000567890ABCD",
        (0, 4, 37): "567890AB-C",
        # R5: *法定代表人, 单位电话, 邮政编码
        (0, 5, 7): "Liu Yang",
        (0, 5, 24): "021-62345678",
        (0, 5, 37): "200120",
        # R6: 通信地址, 单位E-MAIL地址
        (0, 6, 7): "上海市浦东新区陆家嘴环路1000号恒生银行大厦28层",
        (0, 6, 30): "liuyang@starchentrade.com",
        # R7: *市场细分 (勾选中小企业版)
        (0, 7, 7): "☑中小企业版",
        # R8: 服务选择 WEB渠道
        (0, 8, 6): "☑账户查询  ☑转账汇划  ☑对账服务",
        # R10: 所属集团名称 (无集团, 留空)

        # 客户账户信息 — R14: 第一行数据
        # R12/R13 are headers, R14 is first data row
        # C2-4=序号, C5-14=*账号, C15-16=*币种, C17-19=*账户类型, C20-28=*账户名称, C29-32=*服务代码, C33-34=*授权模板, C35-38=单笔限额, C39-41=每日累计限额
        (0, 14, 2): "1",
        (0, 14, 5): "454676800100123456",
        (0, 14, 15): "CNY",
        (0, 14, 20): "上海星辰国际贸易有限公司",
        (0, 14, 35): "500,000",
        (0, 14, 39): "2,000,000",
        # R15: 第二个账户
        (0, 15, 2): "2",
        (0, 15, 5): "454676800100789012",
        (0, 15, 15): "USD",
        (0, 15, 20): "上海星辰国际贸易有限公司",
        (0, 15, 35): "100,000",
        (0, 15, 39): "500,000",

        # 操作员信息 — R21: 第一行数据 (R19/R20 are headers)
        # C1-7=*姓名, C8-12=*角色/功能, C13-15=*可操作账户序号, C16-17=*证件类型, C18-22=*证件号码, C23-29=*密码发送, C30-35=*移动电话
        (0, 21, 1): "Liu Yang",
        (0, 21, 8): "制单员",
        (0, 21, 13): "1,2",
        (0, 21, 16): "身份证",
        (0, 21, 18): "310101199001011234",
        (0, 21, 23): "☑密码信封",
        (0, 21, 30): "13912345678",
        # R22: 第二操作员
        (0, 22, 1): "Wang Fang",
        (0, 22, 8): "复核员",
        (0, 22, 13): "1,2",
        (0, 22, 16): "身份证",
        (0, 22, 18): "310101199205051234",
        (0, 22, 23): "☑手机短信",
        (0, 22, 30): "13698765432",
    },
    eflow_data={
        "flow_id": "EF2026041401",
        "company": {
            "name": "上海星辰国际贸易有限公司",
            "cert_type": "统一社会信用代码",
            "cert_number": "91310000567890ABCD"
        },
        "account": {
            "bank_name": "中国银行",
            "branch": "上海浦东支行",
            "account_number": "454676800100123456"
        },
        "operator": {
            "name": "Liu Yang",
            "id_type": "身份证",
            "id_number": "310101199001011234"
        },
        "handler": {"name": "Liu Yang", "id_type": "身份证", "id_number": "310101199001011234"},
        "activity": "开通企业网上银行",
        "permissions": {"level": "Level_A", "single_limit": 500000, "daily_limit": 2000000},
        "apply_date": "2026-04-14"
    },
    id_img_src=ID_IMG_1
)


# ######################################################################
# CASE 021: CCB — 高限额 + 敏感行业 (INFO)
# Template: [perfect]ccb_corp_app.docx — Table 0 (17 rows x 6 cols)
# R0: 单位名称 C0=label, C1-5=value
# R1: 证件类型 C0=label, C1=val; 证件号码 C2-3=label, C4-5=val
# R2: 地址 C0=label, C1-5=val
# R3: 邮政编码 C0=label, C1=val; 法定代表人 C2-3=label, C4-5=val
# R4: 联系人姓名 C0=label, C1=val; 联系电话 C2-3=label, C4-5=val
# R5: E-Mail地址 C0=label, C1=val; 传真号 C2-3=label, C4-5=val
# R6: 客户服务选择 C0=label, C1-5=checkbox (☑网上银行服务)
# R7: 账户签约类型 C0=label, C1-5=checkbox (☑账户签约)
# R8: 账户名称 C0=label, C1=val; 开户行 C2-3=label, C4-5=val
# R9: 新账号 C0=label, C1=val; 旧账号 C2-3=label, C4-5=val
# ######################################################################
print("\n=== Generating case_021_ccb_high_limit ===")
create_case(
    case_id="case_021_ccb_high_limit",
    description="[风险] CCB建行 - 单笔限额超500万 + 珠宝贸易敏感行业 (INFO)",
    template_src=CCB,
    field_map={
        (0, 0, 1): "深圳市鑫达珠宝贸易有限公司",                      # 单位名称
        (0, 1, 1): "统一社会信用代码",                                  # 证件类型
        (0, 1, 4): "91440300MA5F123456",                                # 证件号码
        (0, 2, 1): "深圳市罗湖区翠竹路水贝珠宝大厦12层",               # 地址
        (0, 3, 1): "518000",                                            # 邮政编码
        (0, 3, 4): "赵强",                                              # 法定代表人
        (0, 4, 1): "赵强",                                              # 联系人姓名
        (0, 4, 4): "0755-22001234",                                     # 联系电话
        (0, 5, 1): "zhao@xindajewelry.com",                             # E-Mail地址
        (0, 5, 4): "0755-22001235",                                     # 传真号
        (0, 6, 1): "☑网上银行服务",                                     # 客户服务选择
        (0, 7, 1): "☑账户签约",                                         # 账户签约类型
        (0, 8, 1): "深圳市鑫达珠宝贸易有限公司",                       # 账户名称
        (0, 8, 4): "建设银行深圳罗湖支行",                             # 开户行
        (0, 9, 1): "4420 1234 5678 0001",                               # 新账号
    },
    eflow_data={
        "flow_id": "EF2026042101",
        "company": {"name": "深圳市鑫达珠宝贸易有限公司", "cert_type": "统一社会信用代码", "cert_number": "91440300MA5F123456"},
        "account": {"bank_name": "中国建设银行", "branch": "深圳罗湖支行", "account_number": "4420123456780001"},
        "operator": {"name": "赵强", "id_type": "身份证", "id_number": "440301198806061234"},
        "handler": {"name": "赵强", "id_type": "身份证", "id_number": "440301198806061234"},
        "activity": "珠宝黄金贸易网银开通",
        "permissions": {"level": "Level_S", "single_limit": 10000000, "daily_limit": 50000000},
        "apply_date": "2026-04-21"
    },
    id_img_src=ID_IMG_1
)


# ######################################################################
# CASE 012: BOC HK — PASS (Hong Kong complex merged table)
# Template: 中国银行香港表1.docx — Table 0 (38 rows x 24 cols)
# R3: *申请单位名称 C0-6=label, C7-23=value
# R4: 证件类别 C0-3=label, C4-11=checkbox, C12-13=证件号码, C14-23=value
# R5: 法定代表人 C0-3=label, C4-11=value, C12-13=单位电话, C14-23=value
# R6: 通信地址 C0-3=label, C4-13=value, C14-18=单位E-MAIL地址, C19-23=value
# R12: 客户账户信息 第一行数据: C1-8=*账号, C9-10=*币种, C11-15=*账户名称, C16-20=服务代码, C21-22=单笔限额, C23=每日累计限额
# R22: 操作员信息 第一行数据: C1-5=*姓名, C6-8=*证件类型, C9-14=*证件号码, C15=*密码发送, C16-20=*移动电话
# ######################################################################
print("\n=== Generating case_012_boc_hk_pass ===")
create_case(
    case_id="case_012_boc_hk_pass",
    description="[正例] 香港中行 - 复杂合并单元格(38x24表格) 全面填写，信息一致",
    template_src=BOC_HK,
    field_map={
        # 服务类型: 勾选"新增" -> "网银新开户"
        (0, 0, 4): "☑网银新开户（包括：客户注册、服务维护、账户注册、操作员注册、认证工具申请/绑定等一系列新增服务类型）",
        # R3: *申请单位名称
        (0, 3, 7): "Hong Kong Bright Future Trading Ltd",
        # R4: 证件类别(checkbox) + 证件号码
        (0, 4, 4): "☑营业执照",
        (0, 4, 14): "91110000MA0EXAMPLE",
        # R5: 法定代表人 + 单位电话
        (0, 5, 4): "Chen Ming",
        (0, 5, 14): "+852 2888 1234",
        # R6: 通信地址 + 单位E-MAIL地址
        (0, 6, 4): "Unit 2301, Tower 1, Lippo Centre, Admiralty, HK",
        (0, 6, 19): "chenming@brightfuture.hk",
        # R8: 服务选择 WEB渠道
        (0, 8, 2): "☑IA海外企业账户管理  ☑IB海外企业转账汇划",
        # R12: 第一行账户 — 账号, 币种, 名称, 单笔限额, 累计限额
        (0, 12, 1): "012-676-0-012345-6",
        (0, 12, 9): "HKD",
        (0, 12, 11): "Hong Kong Bright Future Trading Ltd",
        (0, 12, 21): "500,000",
        (0, 12, 23): "2,000,000",
        # R13: 第二行账户
        (0, 13, 1): "012-676-0-012345-7",
        (0, 13, 9): "USD",
        (0, 13, 11): "Hong Kong Bright Future Trading Ltd",
        (0, 13, 21): "200,000",
        (0, 13, 23): "1,000,000",
        # R22: 操作员1
        (0, 22, 1): "Chen Ming",
        (0, 22, 6): "Passport",
        (0, 22, 9): "K12345678",
        (0, 22, 15): "密码信封",
        (0, 22, 16): "+852 9123 4567",
        # R23: 操作员2
        (0, 23, 1): "Li Na",
        (0, 23, 6): "身份证",
        (0, 23, 9): "440301199501011234",
        (0, 23, 15): "手机短信",
        (0, 23, 16): "+852 9888 7654",
    },
    eflow_data={
        "flow_id": "EF2026041201",
        "company": {"name": "Hong Kong Bright Future Trading Ltd", "cert_type": "Business Registration Certificate", "cert_number": "91110000MA0EXAMPLE"},
        "account": {"bank_name": "Bank of China (Hong Kong)", "branch": "Admiralty Branch", "account_number": "012-676-0-012345-6"},
        "operator": {"name": "Chen Ming", "id_type": "Passport", "id_number": "K12345678"},
        "handler": {"name": "Chen Ming", "id_type": "Passport", "id_number": "K12345678"},
        "activity": "网银新开户",
        "permissions": {"level": "Level_A", "single_limit": 500000, "daily_limit": 2000000},
        "apply_date": "2026-04-12"
    },
    id_img_src=ID_IMG_1
)


# ######################################################################
# CASE 020: BOC HK — Operator name different (WARNING)
# ######################################################################
print("\n=== Generating case_020_boc_hk_operator_diff ===")
create_case(
    case_id="case_020_boc_hk_operator_diff",
    description="[负例] 香港中行 - eflow经办人(Chen Ming)与文档操作员(Zhang Fang)姓名不同 (WARNING)",
    template_src=BOC_HK,
    field_map={
        (0, 0, 4): "☑网银新开户",
        (0, 3, 7): "Hong Kong Bright Future Trading Ltd",
        (0, 4, 4): "☑营业执照",
        (0, 4, 14): "91110000MA0EXAMPLE",
        (0, 5, 4): "Chen Ming",
        (0, 5, 14): "+852 2888 1234",
        (0, 6, 4): "Unit 2301, Tower 1, Lippo Centre, Admiralty, HK",
        (0, 12, 1): "012-676-0-012345-6",
        (0, 12, 9): "HKD",
        (0, 12, 11): "Hong Kong Bright Future Trading Ltd",
        (0, 12, 21): "500,000",
        (0, 12, 23): "2,000,000",
        # 操作员: 填写 Zhang Fang, 但 eflow 里是 Chen Ming -> WARNING
        (0, 22, 1): "Zhang Fang",
        (0, 22, 6): "ID Card",
        (0, 22, 9): "K88888888",
        (0, 22, 15): "密码信封",
        (0, 22, 16): "+852 9111 2222",
    },
    eflow_data={
        "flow_id": "EF2026042001",
        "company": {"name": "Hong Kong Bright Future Trading Ltd", "cert_type": "Business Registration Certificate", "cert_number": "91110000MA0EXAMPLE"},
        "account": {"bank_name": "Bank of China (Hong Kong)", "branch": "Admiralty Branch", "account_number": "012-676-0-012345-6"},
        "operator": {"name": "Chen Ming", "id_type": "Passport", "id_number": "K12345678"},
        "handler": {"name": "Chen Ming", "id_type": "Passport", "id_number": "K12345678"},
        "activity": "网银新开户",
        "permissions": {"level": "Level_A", "single_limit": 500000, "daily_limit": 2000000},
        "apply_date": "2026-04-20"
    },
    id_img_src=ID_IMG_2
)


# ######################################################################
# CASE 013: ICBC Cert — PASS
# Template: ICBC企业客户证书 — Table 1 (25x13) 客户基本信息
# R1: 单位中文名称 C0-1=label, C2-9=value, C10=单位英文名称, C11-12=value
# R2: 单位地址 C0-1=label, C2-12=value (包含邮政编码)
# R3: 法人代表 C0-1=label, C2-9=value, C10=联系电话, C11-12=value
# R4: 传真电话 C0-1=label, C2-9=value, C10=E-MAIL地址, C11-12=value
# R5: 主申请账户账号 C0-1=label, C2-9=value, C10=行业分类, C11-12=value
# R6: 缴费账户账号 C0-1=label, C2-12=value
# ######################################################################
print("\n=== Generating case_013_icbc_cert_pass ===")
create_case(
    case_id="case_013_icbc_cert_pass",
    description="[正例] ICBC工行证书及分支机构信息表 - 全面填写，信息一致",
    template_src=ICBC,
    field_map={
        # --- Table 1: 客户基本信息 ---
        (1, 1, 2): "北京中远国际物流有限公司",      # 单位中文名称
        (1, 1, 11): "Beijing Zhongyuan Intl Logistics",  # 单位英文名称
        (1, 2, 2): "北京市朝阳区建国门外大街甲6号中环世贸D座18层  邮政编码：100022",  # 地址+邮编
        (1, 3, 2): "Wang Jun",                             # 法人代表
        (1, 3, 11): "010-65001234",                          # 联系电话
        (1, 4, 2): "010-65001235",                           # 传真电话
        (1, 4, 11): "wangjun@zhongyuanlogistics.com",        # E-MAIL
        (1, 5, 2): "0200004509201234567",                    # 主申请账户账号
        (1, 5, 11): "交通运输",                               # 行业分类
        (1, 6, 2): "0200004509201234567  (开户行所在地：北京)",  # 缴费账户账号

        # --- Table 0: 批量扣费业务信息 ---
        # R5: 收款账号(已预填) + 单笔最大金额
        (0, 5, 11): "50,000,000.00",
        # R9: 第一行缴费企业
        (0, 9, 1): "北京中远国际物流有限公司",
        (0, 9, 4): "0200004509201234567",
        (0, 9, 8): "工行朝阳支行",
        (0, 9, 12): "增加☑",
    },
    eflow_data={
        "flow_id": "EF2026041301",
        "company": {"name": "北京中远国际物流有限公司", "cert_type": "统一社会信用代码", "cert_number": "91110000351234567X"},
        "account": {"bank_name": "中国工商银行", "branch": "北京朝阳支行", "account_number": "0200004509201234567"},
        "operator": {"name": "Wang Jun", "id_type": "身份证", "id_number": "110101198501012345"},
        "handler": {"name": "Wang Jun", "id_type": "身份证", "id_number": "110101198501012345"},
        "activity": "企业网银证书申请",
        "permissions": {"level": "Level_B", "single_limit": 1000000, "daily_limit": 5000000},
        "apply_date": "2026-04-13"
    },
    id_img_src=ID_IMG_2
)


# ######################################################################
# CASE 022: BOC France — Expired ID (CRITICAL)
# ######################################################################
print("\n=== Generating case_022_boc_france_expired_id ===")
create_case(
    case_id="case_022_boc_france_expired_id",
    description="[负例] BOC法国 - 经办人证件已过期(expiry_date=2024-01-15) (CRITICAL)",
    template_src=BOC_FRANCE,
    field_map={
        (2, 0, 1): "Europe Tech Innovation SAS",
        (2, 1, 1): "990088",
        (2, 2, 1): "8 Avenue des Champs-Elysees, 75008 Paris",
        (2, 3, 1): "admin@europetech.fr",
        (2, 4, 1): "+33 1 56 89 0001",
        (2, 4, 3): "+33 1 56 89 0002",
        (3, 6, 0): "FR7630004000039876543210187",
        (3, 6, 4): "EUR",
        (3, 6, 5): "100,000.00",
        (3, 6, 7): "500,000.00",
        (8, 1, 1): "Zhang San",
        (8, 2, 1): "Europe Tech Innovation SAS",
        (8, 3, 1): "Passport",
        (8, 3, 3): "G55667788",
        (8, 4, 1): "zhangsan@europetech.fr",
        (8, 5, 3): "+33 1 56 89 0001",
        (8, 6, 1): "+33 6 98 76 54 32",
    },
    eflow_data={
        "flow_id": "EF2026042201",
        "company": {"name": "Europe Tech Innovation SAS", "cert_type": "Business Registration", "cert_number": "RCS Paris B 990 088 777"},
        "account": {"bank_name": "Bank of China Paris Branch", "branch": "Paris Main", "account_number": "FR7630004000039876543210187"},
        "operator": {"name": "Zhang San", "id_type": "Passport", "id_number": "G55667788", "expiry_date": "2024-01-15"},
        "handler": {"name": "Zhang San", "id_type": "Passport", "id_number": "G55667788"},
        "activity": "Apply for e-Corp Internet Banking",
        "permissions": {"level": "Level_A", "single_limit": 100000, "daily_limit": 500000},
        "apply_date": "2026-04-22"
    },
    id_img_src=ID_IMG_2
)


# ======================================================================
# Summary
# ======================================================================
print("\n" + "="*60)
print("ALL CASES RE-GENERATED!")
print("="*60)
new_cases = [d.name for d in sorted(TEST_DATA_DIR.iterdir()) if d.is_dir() and d.name.startswith("case_0")]
print(f"Total test cases: {len(new_cases)}")
for c in new_cases:
    files = [f.name for f in (TEST_DATA_DIR / c).iterdir() if f.is_file()]
    print(f"  - {c}: {files}")
