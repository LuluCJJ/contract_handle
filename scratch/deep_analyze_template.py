# -*- coding: utf-8 -*-
"""
Deep analysis: Print every cell of every table with exact (table, row, col) coordinates.
Focus on understanding label-value relationships.
"""
import sys, os
sys.path.append(os.getcwd())
from docx import Document
from pathlib import Path

def deep_analyze(docx_path, name):
    doc = Document(docx_path)
    print(f"\n{'='*80}")
    print(f"TEMPLATE: {name}")
    print(f"Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")
    print(f"{'='*80}")
    
    for ti, table in enumerate(doc.tables):
        print(f"\n  === Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols ===")
        
        # Track which cells are merged by checking if cell objects repeat
        for ri, row in enumerate(table.rows):
            seen_cells = {}  # id(cell) -> first col index
            logical_cells = []  # (start_col, end_col, text)
            
            for ci, cell in enumerate(row.cells):
                cell_id = id(cell)
                if cell_id in seen_cells:
                    # This is a merged continuation, update end_col
                    for lc in logical_cells:
                        if lc[0] == seen_cells[cell_id]:
                            lc[2] = ci  # update end col
                    continue
                seen_cells[cell_id] = ci
                text = cell.text.strip().replace('\n', ' | ')[:60]
                logical_cells.append([ci, ci, ci, text])  # [start_col, end_col_unused, end_col, text]
            
            # Print logical cells
            parts = []
            for lc in logical_cells:
                start_col = lc[0]
                end_col = lc[2]
                text = lc[3]
                if start_col == end_col:
                    parts.append(f"  C{start_col}: [{text}]")
                else:
                    parts.append(f"  C{start_col}-{end_col}: [{text}]")
            
            print(f"  R{ri}: " + " | ".join(parts))

# Analyze key templates
CONVERTED = Path(r"d:\AI\project\contract_handle\_converted_docx")
TEMPLATE = Path(r"d:\AI\project\contract_handle\inputs\bank_template")

# 1. The BOC domestic form the user showed
# 2. CCB template
print("\n" + "#"*80)
print("# CCB - [perfect]ccb_corp_app")
print("#"*80)
deep_analyze(str(CONVERTED / "[perfect]ccb_corp_app.docx"), "CCB Corp App")

# 3. BOC France template
print("\n" + "#"*80)
print("# BOC France - [1]boc_corp_app")
print("#"*80)
deep_analyze(str(CONVERTED / "[1]boc_corp_app.docx"), "BOC France")

# 4. ICBC template
print("\n" + "#"*80)
print("# ICBC - 企业客户证书及分支机构信息表")
print("#"*80)
deep_analyze(str(CONVERTED / "ICBC企业客户证书及分支机构信息表.docx"), "ICBC Cert")

# 5. BOC HK Table 1
print("\n" + "#"*80)
print("# BOC HK - 中国银行香港表1")
print("#"*80)
deep_analyze(str(TEMPLATE / "中国银行香港表1.docx"), "BOC HK Table 1")

