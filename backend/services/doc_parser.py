"""
文档解析模块 — 将 .doc/.docx/.pdf 解析为结构化文本
V2.0 - 增加 PDF(可编辑)解析 & 合并单元格去重
"""
import os
import tempfile
from pathlib import Path
from docx import Document


def convert_doc_to_docx(doc_path: str) -> str:
    """
    将 .doc 文件通过 Word COM 自动化转换为 .docx
    返回临时 .docx 文件路径
    """
    import win32com.client as win32

    word = win32.gencache.EnsureDispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0

    try:
        abs_path = os.path.abspath(doc_path)
        tmp_dir = tempfile.mkdtemp(prefix="contract_handle_")
        basename = Path(doc_path).stem + ".docx"
        dst = os.path.join(tmp_dir, basename)

        doc = word.Documents.Open(abs_path, ReadOnly=True)
        doc.SaveAs2(os.path.abspath(dst), FileFormat=12)  # wdFormatXMLDocument
        doc.Close(False)
        return dst
    finally:
        word.Quit()


def _dedup_row(cells: list[str]) -> list[str]:
    """
    对因合并单元格导致的连续重复列进行去重。
    例如 ['A','A','A','B','B','C'] -> ['A','B','C']
    """
    if not cells:
        return cells
    deduped = [cells[0]]
    for c in cells[1:]:
        if c != deduped[-1]:
            deduped.append(c)
    return deduped


def parse_docx(docx_path: str) -> dict:
    """
    解析 .docx 文件，返回结构化内容。
    输出包含段落文本和表格（Markdown 格式，已去重合并单元格）。
    """
    doc = Document(docx_path)

    # 提取非空段落
    paragraphs = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append(text)

    # 提取表格为 Markdown 格式（带合并单元格去重）
    tables_md = []
    tables_raw = []  # 原始行列数据

    for ti, table in enumerate(doc.tables):
        rows_data = []
        md_lines = []

        for ri, row in enumerate(table.rows):
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace("\n", " | ")
                cells.append(cell_text)
            
            # 去重合并单元格
            deduped = _dedup_row(cells)
            rows_data.append(deduped)

            # Markdown 表格行
            md_line = "| " + " | ".join(deduped) + " |"
            md_lines.append(md_line)

            # 加表头分隔线
            if ri == 0:
                sep = "| " + " | ".join(["---"] * len(deduped)) + " |"
                md_lines.append(sep)

        tables_md.append("\n".join(md_lines))
        tables_raw.append(rows_data)

    return {
        "paragraphs": paragraphs,
        "tables_markdown": tables_md,
        "tables_raw": tables_raw,
        "num_paragraphs": len(paragraphs),
        "num_tables": len(tables_md),
    }


def parse_pdf(pdf_path: str) -> dict:
    """
    解析可编辑 PDF 文件。
    使用 PyMuPDF 提取文本段落和表格。
    """
    import fitz  # PyMuPDF

    doc = fitz.open(pdf_path)
    paragraphs = []
    tables_md = []
    tables_raw = []

    for page_num, page in enumerate(doc):
        # 1. 提取段落文本
        text = page.get_text("text")
        for line in text.split("\n"):
            line = line.strip()
            if line:
                paragraphs.append(line)

        # 2. 提取表格（PyMuPDF 内置表格探测）
        try:
            page_tables = page.find_tables()
            for tab in page_tables:
                rows_data = []
                md_lines = []
                extracted = tab.extract()
                for ri, row in enumerate(extracted):
                    cells = [(c or "").strip() for c in row]
                    deduped = _dedup_row(cells)
                    rows_data.append(deduped)

                    md_line = "| " + " | ".join(deduped) + " |"
                    md_lines.append(md_line)

                    if ri == 0:
                        sep = "| " + " | ".join(["---"] * len(deduped)) + " |"
                        md_lines.append(sep)

                if rows_data:
                    tables_md.append("\n".join(md_lines))
                    tables_raw.append(rows_data)
        except Exception as e:
            print(f"[DocParser] PDF table extraction failed on page {page_num}: {e}")

    doc.close()

    return {
        "paragraphs": paragraphs,
        "tables_markdown": tables_md,
        "tables_raw": tables_raw,
        "num_paragraphs": len(paragraphs),
        "num_tables": len(tables_md),
    }


def parse_document(file_path: str) -> dict:
    """
    统一入口：根据扩展名自动处理 .doc / .docx / .pdf
    """
    ext = Path(file_path).suffix.lower()

    if ext == ".docx":
        return parse_docx(file_path)
    elif ext == ".doc":
        converted = convert_doc_to_docx(file_path)
        try:
            return parse_docx(converted)
        finally:
            # 清理临时文件
            try:
                os.unlink(converted)
                os.rmdir(os.path.dirname(converted))
            except OSError:
                pass
    elif ext == ".pdf":
        return parse_pdf(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}，仅支持 .doc / .docx / .pdf")


def get_full_text_for_llm(parsed: dict) -> str:
    """
    将解析结果组装成一段完整文本，供 LLM 分析。
    段落 + 所有表格的 Markdown 表示。
    """
    parts = []

    if parsed["paragraphs"]:
        parts.append("## 文档段落内容\n")
        for p in parsed["paragraphs"]:
            parts.append(p)
        parts.append("")

    for i, tmd in enumerate(parsed["tables_markdown"]):
        parts.append(f"## 表格 {i + 1}\n")
        parts.append(tmd)
        parts.append("")

    return "\n".join(parts)
