"""
批量生成测试用例 — 基于真实证件信息 + 银行模板
生成正例、负例和风险样例
"""
from docx import Document
import json, os, shutil

# === 路径配置 ===
CCB_TEMPLATE = r"d:\AI\project\contract_handle\_converted_docx\[perfect]ccb_corp_app.docx"
BOC_TEMPLATE = r"d:\AI\project\contract_handle\_converted_docx\[1]boc_corp_app.docx"
PASSPORT_DIR = r"d:\AI\project\contract_handle\inputs\passports"
OUTPUT_DIR = r"d:\AI\project\contract_handle\test_data"

log_lines = []
def log(msg): log_lines.append(str(msg))

# === 填充函数 ===

def fill_ccb(data, output_path):
    """填充建行模板"""
    doc = Document(CCB_TEMPLATE)
    table = doc.tables[0]
    field_map = {
        (0, 1): data.get("company_name", ""),
        (1, 1): data.get("cert_type", ""),
        (1, 4): data.get("cert_number", ""),
        (2, 1): data.get("address", ""),
        (3, 1): data.get("postal_code", ""),
        (3, 4): data.get("legal_rep", ""),
        (4, 1): data.get("contact_name", ""),
        (4, 4): data.get("contact_phone", ""),
        (5, 1): data.get("email", ""),
        (5, 4): data.get("fax", ""),
        (8, 1): data.get("account_name", ""),
        (8, 4): data.get("bank_branch", ""),
        (9, 1): data.get("new_account", ""),
    }
    for (r, c), val in field_map.items():
        if val:
            cell = table.cell(r, c)
            for p in cell.paragraphs: p.text = ""
            cell.paragraphs[0].text = val
    
    # 勾选服务类型
    if data.get("service_check"):
        cell = table.cell(6, 1)
        cell.paragraphs[0].text = data["service_check"]
    if data.get("signup_check"):
        cell = table.cell(7, 1)
        cell.paragraphs[0].text = data["signup_check"]
    
    doc.save(output_path)

def fill_boc(data, output_path):
    """填充中行模板（公司信息表2 + 操作员信息表8）"""
    doc = Document(BOC_TEMPLATE)
    
    # 表2: 公司基本信息
    t2 = doc.tables[2]
    t2_map = {
        (0, 1): data.get("company_name_en", ""),
        (1, 1): data.get("account_prefix", ""),
        (2, 1): data.get("address_en", ""),
        (3, 1): data.get("email", ""),
        (4, 1): data.get("phone", ""),
        (4, 3): data.get("fax", ""),
    }
    for (r, c), val in t2_map.items():
        if val:
            cell = t2.cell(r, c)
            for p in cell.paragraphs: p.text = ""
            cell.paragraphs[0].text = val
    
    # 表8: 操作员信息
    t8 = doc.tables[8]
    t8_map = {
        (1, 1): data.get("operator_name", ""),
        (2, 1): data.get("company_name_en", ""),
        (3, 1): data.get("operator_id_type", ""),
        (3, 3): data.get("operator_id_number", ""),
        (4, 1): data.get("operator_email", ""),
        (5, 1): data.get("operator_fax", ""),
        (5, 3): data.get("operator_phone", ""),
        (6, 1): data.get("operator_mobile", ""),
    }
    for (r, c), val in t8_map.items():
        if val:
            cell = t8.cell(r, c)
            for p in cell.paragraphs: p.text = ""
            cell.paragraphs[0].text = val
    
    doc.save(output_path)

def make_eflow(data):
    """生成 E-Flow JSON"""
    return {
        "flow_id": data.get("flow_id", ""),
        "company": {
            "name": data.get("company_name", ""),
            "name_en": data.get("company_name_en", ""),
            "cert_type": data.get("company_cert_type", ""),
            "cert_number": data.get("company_cert_number", "")
        },
        "account": {
            "bank_name": data.get("bank_name", ""),
            "branch": data.get("bank_branch", ""),
            "account_number": data.get("account_number", "")
        },
        "operator": {
            "name": data.get("operator_name", ""),
            "id_type": data.get("operator_id_type", ""),
            "id_number": data.get("operator_id_number", "")
        },
        "handler": {
            "name": data.get("handler_name", ""),
            "id_type": data.get("handler_id_type", ""),
            "id_number": data.get("handler_id_number", "")
        },
        "activity": data.get("activity", ""),
        "permissions": {
            "level": data.get("perm_level", ""),
            "single_limit": data.get("single_limit", 0),
            "daily_limit": data.get("daily_limit", 0)
        },
        "apply_date": "2026-04-03"
    }

def gen_case(case_id, description, template_type, word_data, eflow_data, passport_file):
    """生成一套完整测试用例"""
    case_dir = os.path.join(OUTPUT_DIR, case_id)
    os.makedirs(case_dir, exist_ok=True)
    
    # 1. 填充 Word
    word_path = os.path.join(case_dir, "bank_app.docx")
    if template_type == "ccb":
        fill_ccb(word_data, word_path)
    else:
        fill_boc(word_data, word_path)
    
    # 2. 生成 E-Flow JSON
    eflow = make_eflow(eflow_data)
    with open(os.path.join(case_dir, "eflow.json"), "w", encoding="utf-8") as f:
        json.dump(eflow, f, ensure_ascii=False, indent=2)
    
    # 3. 复制证件图片
    if passport_file:
        src = os.path.join(PASSPORT_DIR, passport_file)
        if os.path.exists(src):
            ext = os.path.splitext(passport_file)[1]
            shutil.copy2(src, os.path.join(case_dir, f"id_document{ext}"))
    
    # 4. 写说明文件
    with open(os.path.join(case_dir, "README.md"), "w", encoding="utf-8") as f:
        f.write(f"# {case_id}\n\n**场景**: {description}\n")
    
    log(f"✅ {case_id}: {description}")


# ====================================================================
# 测试用例定义
# ====================================================================

# --- Case 001: 正例 - CCB 全部一致 ---
gen_case(
    case_id="case_001_pass",
    description="[正例] CCB建行 - 所有信息完全一致，预期通过",
    template_type="ccb",
    word_data={
        "company_name": "深圳市前海科技有限公司",
        "cert_type": "统一社会信用代码",
        "cert_number": "91440300MA5EXXXX1L",
        "address": "深圳市南山区科技园南区",
        "postal_code": "518000",
        "legal_rep": "张光",
        "contact_name": "张光",
        "contact_phone": "0755-86001234",
        "email": "zhangguang@qianhai-tech.com",
        "fax": "0755-86001235",
        "account_name": "深圳市前海科技有限公司",
        "bank_branch": "建设银行深圳南山支行",
        "new_account": "4420156430005200",
        "service_check": "☑网上银行服务    □电子对账服务       □注销服务",
        "signup_check": "☑账户签约        □  账户更改         □账户注销",
    },
    eflow_data={
        "flow_id": "EF2026040301",
        "company_name": "深圳市前海科技有限公司",
        "company_name_en": "Qianhai Technology Co., Limited",
        "company_cert_type": "统一社会信用代码",
        "company_cert_number": "91440300MA5EXXXX1L",
        "bank_name": "中国建设银行",
        "bank_branch": "深圳南山支行",
        "account_number": "4420156430005200",
        "operator_name": "张光",
        "operator_id_type": "身份证",
        "operator_id_number": "1324331974",  # 匹配身份证1
        "handler_name": "张光",
        "handler_id_type": "身份证",
        "handler_id_number": "1324331974",
        "activity": "开通网上银行",
        "perm_level": "Level_A",
        "single_limit": 500000,
        "daily_limit": 1000000,
    },
    passport_file="身份证1.jpg",
)

# --- Case 002: 正例 - BOC 全部一致 ---
gen_case(
    case_id="case_002_pass",
    description="[正例] BOC中行 - 所有信息完全一致，预期通过",
    template_type="boc",
    word_data={
        "company_name_en": "Qianhai Technology Co., Limited",
        "account_prefix": "C12345",
        "address_en": "No.88 Keyuan Road, Nanshan, Shenzhen",
        "email": "info@qianhai-tech.com",
        "phone": "+86-755-86001234",
        "fax": "+86-755-86001235",
        "operator_name": "SANTA CLAUS",
        "operator_id_type": "Passport",
        "operator_id_number": "N1234567",
        "operator_email": "santa@qianhai-tech.com",
        "operator_fax": "+86-755-86001235",
        "operator_phone": "+86-13800001234",
        "operator_mobile": "+86-13800001234",
    },
    eflow_data={
        "flow_id": "EF2026040302",
        "company_name": "前海科技有限公司",
        "company_name_en": "Qianhai Technology Co., Limited",
        "company_cert_type": "USCI",
        "company_cert_number": "91440300MA5EXXXX1L",
        "bank_name": "Bank of China",
        "bank_branch": "Paris Branch",
        "account_number": "C12345XXXXXXX",
        "operator_name": "SANTA CLAUS",
        "operator_id_type": "Passport",
        "operator_id_number": "N1234567",
        "handler_name": "SANTA CLAUS",
        "handler_id_type": "Passport",
        "handler_id_number": "N1234567",
        "activity": "Open Online Banking",
        "perm_level": "Level_A",
        "single_limit": 500000,
        "daily_limit": 1000000,
    },
    passport_file="护照1.jpg",
)

# --- Case 003: 负例 - 证件号不一致 ---
gen_case(
    case_id="case_003_fail_id",
    description="[负例] CCB建行 - Word证件号与E-Flow/证件不一致 (CRITICAL)",
    template_type="ccb",
    word_data={
        "company_name": "深圳市前海科技有限公司",
        "cert_type": "统一社会信用代码",
        "cert_number": "91440300MA5EXXXX2M",  # ← 故意错！与 E-Flow 不一致
        "address": "深圳市南山区科技园南区",
        "postal_code": "518000",
        "legal_rep": "张光",
        "contact_name": "张光",
        "contact_phone": "0755-86001234",
        "email": "zhangguang@qianhai-tech.com",
        "fax": "0755-86001235",
        "account_name": "深圳市前海科技有限公司",
        "bank_branch": "建设银行深圳南山支行",
        "new_account": "4420156430005200",
        "service_check": "☑网上银行服务    □电子对账服务       □注销服务",
        "signup_check": "☑账户签约        □  账户更改         □账户注销",
    },
    eflow_data={
        "flow_id": "EF2026040303",
        "company_name": "深圳市前海科技有限公司",
        "company_name_en": "Qianhai Technology Co., Limited",
        "company_cert_type": "统一社会信用代码",
        "company_cert_number": "91440300MA5EXXXX1L",  # ← 正确的
        "bank_name": "中国建设银行",
        "bank_branch": "深圳南山支行",
        "account_number": "4420156430005200",
        "operator_name": "张光",
        "operator_id_type": "身份证",
        "operator_id_number": "1324331974",
        "handler_name": "张光",
        "handler_id_type": "身份证",
        "handler_id_number": "1324331974",
        "activity": "开通网上银行",
        "perm_level": "Level_A",
        "single_limit": 500000,
        "daily_limit": 1000000,
    },
    passport_file="身份证1.jpg",
)

# --- Case 004: 负例 - 操作员姓名不一致 ---
gen_case(
    case_id="case_004_fail_name",
    description="[负例] BOC中行 - 操作员姓名Word写CLAIRE BENNETT，E-Flow/护照是SANTA CLAUS (CRITICAL)",
    template_type="boc",
    word_data={
        "company_name_en": "Qianhai Technology Co., Limited",
        "account_prefix": "C12345",
        "address_en": "No.88 Keyuan Road, Nanshan, Shenzhen",
        "email": "info@qianhai-tech.com",
        "phone": "+86-755-86001234",
        "fax": "+86-755-86001235",
        "operator_name": "CLAIRE BENNETT",  # ← 故意写错人！
        "operator_id_type": "Passport",
        "operator_id_number": "N1234567",     # 证件号一致
        "operator_email": "claire@qianhai-tech.com",
        "operator_fax": "+86-755-86001235",
        "operator_phone": "+86-13800005678",
        "operator_mobile": "+86-13800005678",
    },
    eflow_data={
        "flow_id": "EF2026040304",
        "company_name": "前海科技有限公司",
        "company_name_en": "Qianhai Technology Co., Limited",
        "company_cert_type": "USCI",
        "company_cert_number": "91440300MA5EXXXX1L",
        "bank_name": "Bank of China",
        "bank_branch": "Paris Branch",
        "account_number": "C12345XXXXXXX",
        "operator_name": "SANTA CLAUS",       # ← 正确的
        "operator_id_type": "Passport",
        "operator_id_number": "N1234567",
        "handler_name": "SANTA CLAUS",
        "handler_id_type": "Passport",
        "handler_id_number": "N1234567",
        "activity": "Open Online Banking",
        "perm_level": "Level_A",
        "single_limit": 500000,
        "daily_limit": 1000000,
    },
    passport_file="护照1.jpg",   # 护照上写的 SANTA CLAUS
)

# --- Case 005: 风险样例 - 经办活动矛盾 ---
gen_case(
    case_id="case_005_risk_activity",
    description="[风险] CCB建行 - E-Flow申请开通网银，但Word勾选了注销服务 (WARNING)",
    template_type="ccb",
    word_data={
        "company_name": "深圳市前海科技有限公司",
        "cert_type": "统一社会信用代码",
        "cert_number": "91440300MA5EXXXX1L",
        "address": "深圳市南山区科技园南区",
        "postal_code": "518000",
        "legal_rep": "张光",
        "contact_name": "张光",
        "contact_phone": "0755-86001234",
        "email": "zhangguang@qianhai-tech.com",
        "fax": "0755-86001235",
        "account_name": "深圳市前海科技有限公司",
        "bank_branch": "建设银行深圳南山支行",
        "new_account": "4420156430005200",
        "service_check": "□网上银行服务    □电子对账服务       ☑注销服务",  # ← 勾了注销！
        "signup_check": "□账户签约        □  账户更改         ☑账户注销",  # ← 勾了注销！
    },
    eflow_data={
        "flow_id": "EF2026040305",
        "company_name": "深圳市前海科技有限公司",
        "company_name_en": "Qianhai Technology Co., Limited",
        "company_cert_type": "统一社会信用代码",
        "company_cert_number": "91440300MA5EXXXX1L",
        "bank_name": "中国建设银行",
        "bank_branch": "深圳南山支行",
        "account_number": "4420156430005200",
        "operator_name": "张光",
        "operator_id_type": "身份证",
        "operator_id_number": "1324331974",
        "handler_name": "张光",
        "handler_id_type": "身份证",
        "handler_id_number": "1324331974",
        "activity": "开通网上银行",      # ← E-Flow 说的是开通
        "perm_level": "Level_A",
        "single_limit": 500000,
        "daily_limit": 1000000,
    },
    passport_file="身份证1.jpg",
)

# --- Case 006: 风险样例 - 账号不一致 ---
gen_case(
    case_id="case_006_risk_account",
    description="[风险] BOC中行 - Word账号前缀与E-Flow不一致 (WARNING)",
    template_type="boc",
    word_data={
        "company_name_en": "Qianhai Technology Co., Limited",
        "account_prefix": "C99999",     # ← 与E-Flow的C12345不一致
        "address_en": "No.88 Keyuan Road, Nanshan, Shenzhen",
        "email": "info@qianhai-tech.com",
        "phone": "+86-755-86001234",
        "fax": "+86-755-86001235",
        "operator_name": "SANTA CLAUS",
        "operator_id_type": "Passport",
        "operator_id_number": "N1234567",
        "operator_email": "santa@qianhai-tech.com",
        "operator_fax": "+86-755-86001235",
        "operator_phone": "+86-13800001234",
        "operator_mobile": "+86-13800001234",
    },
    eflow_data={
        "flow_id": "EF2026040306",
        "company_name": "前海科技有限公司",
        "company_name_en": "Qianhai Technology Co., Limited",
        "company_cert_type": "USCI",
        "company_cert_number": "91440300MA5EXXXX1L",
        "bank_name": "Bank of China",
        "bank_branch": "Paris Branch",
        "account_number": "C12345XXXXXXX",    # ← 正确的
        "operator_name": "SANTA CLAUS",
        "operator_id_type": "Passport",
        "operator_id_number": "N1234567",
        "handler_name": "SANTA CLAUS",
        "handler_id_type": "Passport",
        "handler_id_number": "N1234567",
        "activity": "Open Online Banking",
        "perm_level": "Level_A",
        "single_limit": 500000,
        "daily_limit": 1000000,
    },
    passport_file="护照1.jpg",
)

# --- Case 007: 风险样例 - 证件类型不匹配 ---
gen_case(
    case_id="case_007_risk_idtype",
    description="[风险] BOC中行 - E-Flow写身份证但实际提供护照 (WARNING)",
    template_type="boc",
    word_data={
        "company_name_en": "Qianhai Technology Co., Limited",
        "account_prefix": "C12345",
        "address_en": "No.88 Keyuan Road, Nanshan, Shenzhen",
        "email": "info@qianhai-tech.com",
        "phone": "+86-755-86001234",
        "fax": "+86-755-86001235",
        "operator_name": "SANTA CLAUS",
        "operator_id_type": "ID Card",   # ← Word写了身份证
        "operator_id_number": "N1234567",
        "operator_email": "santa@qianhai-tech.com",
        "operator_fax": "+86-755-86001235",
        "operator_phone": "+86-13800001234",
        "operator_mobile": "+86-13800001234",
    },
    eflow_data={
        "flow_id": "EF2026040307",
        "company_name": "前海科技有限公司",
        "company_name_en": "Qianhai Technology Co., Limited",
        "company_cert_type": "USCI",
        "company_cert_number": "91440300MA5EXXXX1L",
        "bank_name": "Bank of China",
        "bank_branch": "Paris Branch",
        "account_number": "C12345XXXXXXX",
        "operator_name": "SANTA CLAUS",
        "operator_id_type": "ID Card",   # ← E-Flow也写身份证
        "operator_id_number": "N1234567",  # 但号码是护照号格式
        "handler_name": "SANTA CLAUS",
        "handler_id_type": "Passport",
        "handler_id_number": "N1234567",
        "activity": "Open Online Banking",
        "perm_level": "Level_A",
        "single_limit": 500000,
        "daily_limit": 1000000,
    },
    passport_file="护照1.jpg",  # ← 实际提供的是护照！
)

# ====================================================================
# 输出结果
# ====================================================================
log("\n" + "=" * 50)
log(f"共生成 7 个测试用例，目录: {OUTPUT_DIR}")
log("=" * 50)

# 写日志
out_log = os.path.join(os.path.dirname(OUTPUT_DIR), "_gen_all_output.txt")
with open(out_log, "w", encoding="utf-8") as f:
    f.write("\n".join(log_lines))
print(f"Output: {out_log}")
